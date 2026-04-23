"""
File reading and writing with format preservation where possible.
Supported: .docx, .doc (text-only), .txt, .md, .markdown, .rtf,
           .csv, .log, .xml, .json, .ini, .cfg, .yaml, .yml, .toml
"""
import os
import re
import copy
import json
from pathlib import Path

# python-docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SUPPORTED_EXTENSIONS = {
    ".docx", ".doc",
    ".txt", ".md", ".markdown",
    ".rtf", ".csv", ".log",
    ".xml", ".json", ".ini",
    ".cfg", ".yaml", ".yml", ".toml",
}

def extension_supported(path: str) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_EXTENSIONS


# ------------------------------------------------------------------ #
#  DOCX                                                               #
# ------------------------------------------------------------------ #

def read_docx(path: str):
    """
    Return (full_text, doc_object).
    full_text is the concatenated paragraph text with paragraph separators.
    doc_object is the python-docx Document for format-preserving write-back.
    """
    doc = Document(path)
    paragraphs_text = []
    for para in doc.paragraphs:
        paragraphs_text.append(para.text)
    # Also handle tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    paragraphs_text.append(para.text)
    return "\n".join(paragraphs_text), doc


def write_docx(doc, findings, output_path: str):
    """
    Write a new DOCX with redacted text. Preserves paragraph structure.
    """
    import re as _re

    def redact_runs(paragraph, findings_for_para):
        """
        Replace text inside runs while preserving run formatting.
        """
        # Merge all runs into a single string, track run boundaries
        full = paragraph.text
        if not full.strip():
            return
        # Apply replacements on the merged string
        replacements = sorted(findings_for_para, key=lambda f: f.start)
        # We rebuild the paragraph text run-by-run using a simple approach:
        # replace in full text then push back into first run, clear others.
        new_text = full
        offset = 0
        for f in replacements:
            if not f.kept:
                continue
            s = f.start + offset
            e = f.end + offset
            new_text = new_text[:s] + f.tag + new_text[e:]
            offset += len(f.tag) - (f.end - f.start)

        if new_text == full:
            return

        # Push modified text back: put everything into the first run
        runs = paragraph.runs
        if not runs:
            paragraph.text = new_text
            return
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""

    # Build a flat list of (para_object, para_text, para_start_in_full_text)
    # We need to match findings (which reference positions in the full text)
    # back to individual paragraphs.
    full_parts = []
    para_objects = []
    cursor = 0
    para_spans = []  # (start, end) in full text

    for para in doc.paragraphs:
        t = para.text
        para_objects.append(para)
        para_spans.append((cursor, cursor + len(t)))
        full_parts.append(t)
        cursor += len(t) + 1  # +1 for the "\n" separator

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text
                    para_objects.append(para)
                    para_spans.append((cursor, cursor + len(t)))
                    full_parts.append(t)
                    cursor += len(t) + 1

    for idx, (para, (pstart, pend)) in enumerate(zip(para_objects, para_spans)):
        para_findings = []
        for f in findings:
            if not f.kept:
                continue
            # Check overlap
            if f.start < pend and f.end > pstart:
                # Translate to para-local offsets
                local_f = copy.copy(f)
                local_f.start = max(f.start, pstart) - pstart
                local_f.end = min(f.end, pend) - pstart
                para_findings.append(local_f)
        if para_findings:
            redact_runs(para, para_findings)

    doc.save(output_path)


# ------------------------------------------------------------------ #
#  Plain text formats                                                 #
# ------------------------------------------------------------------ #

def read_text(path: str) -> str:
    """Read any plain-text file, auto-detect encoding."""
    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return Path(path).read_text(encoding=enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return Path(path).read_bytes().decode("utf-8", errors="replace")


def write_text(original_path: str, new_text: str, output_path: str):
    Path(output_path).write_text(new_text, encoding="utf-8")


# ------------------------------------------------------------------ #
#  Dispatch                                                           #
# ------------------------------------------------------------------ #

def read_file(path: str):
    """
    Returns (text: str, meta: any).
    meta is the python-docx Document for .docx, or None for text files.
    """
    ext = Path(path).suffix.lower()
    if ext == ".docx":
        return read_docx(path)
    else:
        return read_text(path), None


def write_file(original_path: str, findings, meta, output_path: str):
    ext = Path(original_path).suffix.lower()
    if ext == ".docx" and meta is not None:
        write_docx(meta, findings, output_path)
    else:
        from engine import apply_redactions
        text = read_text(original_path)
        new_text = apply_redactions(text, findings)
        write_text(original_path, new_text, output_path)


def suggest_output_path(original_path: str) -> str:
    p = Path(original_path)
    return str(p.parent / (p.stem + ".anonymized" + p.suffix))


def suggest_report_path(original_path: str) -> str:
    p = Path(original_path)
    return str(p.parent / (p.stem + ".audit.json"))
